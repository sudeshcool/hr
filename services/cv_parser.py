"""
cv_parser.py  –  Extract plain text from PDF and DOCX CV files.
"""
import os
import logging

logger = logging.getLogger(__name__)


def extract_text(file_path: str) -> str:
    """Return plain text extracted from a PDF or DOCX file."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return _extract_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        return _extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_path: str) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction failed for {file_path}: {e}")
        return ""


def _extract_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return '\n'.join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {e}")
        return ""
